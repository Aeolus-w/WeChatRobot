from langchain_community.document_loaders import UnstructuredMarkdownLoader, TextLoader
from langchain_community.vectorstores import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
import os
import fitz  # PyMuPDF 用于处理 PDF 文件
import docx  # python-docx 用于处理 DOCX 文件


# 加载 DOCX 文件
def load_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return "\n".join(full_text)


# 加载 PDF 文件
def load_pdf(file_path):
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text


# 加载文件
def load_documents(dir_path):
    docs = []
    for filename in os.listdir(dir_path):
        file_path = os.path.join(dir_path, filename)
        if filename.endswith('.md'):
            loader = UnstructuredMarkdownLoader(file_path)
            docs.extend(loader.load())
        elif filename.endswith('.txt'):
            loader = TextLoader(file_path)
            docs.extend(loader.load())
        elif filename.endswith('.pdf'):
            pdf_text = load_pdf(file_path)
            docs.append(Document(page_content=pdf_text))
        elif filename.endswith('.docx'):
            docx_text = load_docx(file_path)
            docs.append(Document(page_content=docx_text))
    return docs


# 生成向量知识库
def create_vector_db(documents, vector_kb_folder):
    # 分块
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=150)
    split_docs = text_splitter.split_documents(documents)

    # 加载嵌入模型
    embeddings = HuggingFaceEmbeddings(
        model_name='/root/autodl-tmp/sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2')

    # 创建或更新 Chroma 向量数据库，并持久化到 vector_kb 文件夹
    vectordb = Chroma.from_documents(
        documents=split_docs,
        embedding=embeddings,
        persist_directory=vector_kb_folder
    )

    return vectordb


def update_knowledge_db(session_id):
    # 定位当前 session 文件夹路径，并创建 vector_kb 文件夹
    session_folder = f"/root/autodl-tmp/base_knowledge/{session_id}"
    vector_kb_folder = os.path.join(session_folder, "vector_kb")
    os.makedirs(vector_kb_folder, exist_ok=True)

    # 加载文档并生成或更新向量库
    docs = load_documents(session_folder)
    vector_db = create_vector_db(docs, vector_kb_folder)
    return vector_db


# 主程序
if __name__ == "__main__":
    dir_path = '/root/autodl-tmp/base_knowledge'  # 文档目录
    docs = load_documents(dir_path)
    vector_db = create_vector_db(docs)
